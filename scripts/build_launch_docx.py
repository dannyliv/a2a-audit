#!/usr/bin/env python3
"""Build the local-only launch content docx (Substack + LinkedIn + X thread).

Reflects the current state: public repo + live demo, 114-card on-device corpus
(100% unsigned, 77% no-auth), local-model/no-API-key classifier, and the
posture-vs-conformance framing. Output is gitignored (content/).
"""
from __future__ import annotations

from docx import Document
from docx.shared import Pt

OUT = "/Users/danny/Documents/Claude/Projects/A2A Agent Card auditor/a2a-audit/content/a2a-audit-launch.docx"
REPO = "https://github.com/dannyliv/a2a-audit"
DEMO = "https://dannyliv.github.io/a2a-audit/"

doc = Document()
doc.styles["Normal"].font.name = "Calibri"
doc.styles["Normal"].font.size = Pt(11)


def para(text="", bold=False, italic=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    return p


def rule():
    para("-" * 36)


doc.add_heading("a2a-audit launch content", 0)
para(
    "Local working file, not committed. Three platform drafts: Substack article, "
    "LinkedIn post, X thread. Hook options above each. Stats are aggregate from a "
    "114-card on-device scan (2026-05-29). Responsible framing: public cards, "
    "point-in-time, not accusations.",
    italic=True,
)
para(f"Repo: {REPO}   Live demo: {DEMO}", italic=True)

# ---------------------------------------------------------------- SUBSTACK
doc.add_heading("1) Substack article", 1)

doc.add_heading("Hook options", 2)
para("A. I graded the security of 114 live AI agents. Every single one could be impersonated by anyone on the network. Not one was signed.")
para("B. The fastest way to learn how to attack an AI agent is to read the file it publishes about itself. So I built a tool that grades that file, and pointed it at 114 live agents.")
para("C. Valid is not the same as safe. 114 live AI agents passed as working software. On security posture, the average grade was a C and 100 percent were unsigned.")

doc.add_heading("Headline", 2)
para("Your AI agents publish a map to their own front door, and almost none lock it", bold=True)
doc.add_heading("Sub headline", 2)
para("The A2A agent card is the new public attack surface. I graded 114 live agents. The results are not close.", italic=True)

doc.add_heading("Body", 2)
para("Every agent that speaks Google's Agent2Agent protocol publishes a small public file: its agent card. The card says where the agent lives, how to authenticate, and what it can do. It is the agent's business card and its front door. It is also the first thing a serious attacker reads.")
para("I built an open-source tool, a2a-audit, that reads those cards the way an attacker would and turns each one into a single security grade from A to F. Then I pointed it at 114 live agents pulled from a public registry. You can run the same checks yourself in the browser demo, and the full results for all 114 are baked into it.")
para("Across the 114 cards, every one was unsigned. Eighty-eight of them, roughly three quarters, required no authentication at all. The average grade was a C. The few that earned an A were unsigned too. They just had less else wrong.")

para("Valid is not the same as safe", bold=True)
para("There are already good tools that check whether an agent card is valid and works. The official a2a-inspector and a2a-tck do conformance. Cisco's a2a-scanner lists individual issues. None of them answer the question a buyer or a board actually asks: how safe is this agent to trust, in one number I can compare and act on. That is the gap a2a-audit fills. A finding is a fact. A grade is a decision.")

para("An unsigned card has no integrity", bold=True)
para("The unanimous finding, 100 percent unsigned, is the one that should worry a board. The A2A spec defines how to sign a card, using a detached signature over a canonical copy of the card. A signature lets a client confirm the card it fetched is the card the operator actually published. Without one, anyone sitting between the client and the agent, a compromised CDN, a poisoned cache, a stale registry copy, can swap the endpoint or the authentication rules and the client cannot tell. That is a supply-chain weakness on the single file the whole protocol trusts.")

para("Runs on your machine, no API key, nothing leaves your network", bold=True)
para("The skill-description check, which looks for hidden instructions planted in an agent's advertised skills, runs on a local model by default. A small purpose-built classifier ships with the tool; a 7B open-weight model is an option for deeper analysis. Both are permissively licensed and run on a laptop. No cloud key, no data leaving your network. For teams evaluating untrusted agents, that privacy property matters as much as the grade.")

para("The fair objection is that some of these agents are public demos, and a demo with no login is a choice, not a bug. True, which is why undeclared authentication is a medium finding with a caveat, not an automatic failure. But unsigned is different. A demo can be open on purpose. No one ships a card they want anyone in the network path to be able to forge.")
para(f"If you operate an A2A agent, read your own card the way an attacker would, then sign it. If you connect to other agents, grade them first. The tool is open source and the demo is live: {DEMO}")

rule()

# ---------------------------------------------------------------- LINKEDIN
doc.add_heading("2) LinkedIn post", 1)

doc.add_heading("Hook options", 2)
para("A. I graded the security of 114 live AI agents last week. 100 percent were unsigned. 77 percent required no login.")
para("B. Your AI agents publish a file that tells attackers how to reach them. Most teams have never read their own.")
para("C. Valid is not the same as safe. 114 live AI agents worked fine as software. On security posture, the average grade was a C.")

doc.add_heading("Post body", 2)
para("I graded the security posture of 114 live AI agents last week.")
para("Every one of them publishes a small file called an agent card: where it lives, how to authenticate, what it can do. In the Agent2Agent protocol, that card is the unit of trust.")
para("The results were not close.")
para("All 114 were unsigned. 88 of 114, about three quarters, required no authentication. The average security grade was a C.")
para("The unsigned number is the one for the board. A signature lets a client confirm the card it fetched is the card you actually published. Without it, anyone between your client and the agent can swap the endpoint or the auth rules, and no one notices. That is a supply-chain exposure on the one file the protocol trusts.")
para("Conformance tools already tell you a card is valid. They do not tell you it is safe. So I built a2a-audit: it reads the card the way an attacker would, maps each issue to the OWASP Top 10 for Agentic Applications, and returns one grade you can gate in CI. The skill checks run on a local model, so no data leaves your network.")
para("A finding is a fact. A grade is a decision your team can act on.")
para("Two questions if you run agents. Have you read your own card? And what grade would you accept before you let an agent into production?")
para("Live demo and tool in the first comment.")
para("#AIsecurity #AgenticAI #AppSec #CISO #A2A")

doc.add_heading("First comment (links)", 2)
para(f"Live demo (grade any card in the browser): {DEMO}")
para(f"Tool (open source): {REPO}")
para("OWASP Top 10 for Agentic Applications (2026): https://genai.owasp.org/resource/owasp-top-10-for-agentic-applications-for-2026/")

rule()

# ---------------------------------------------------------------- X THREAD
doc.add_heading("3) X thread", 1)

doc.add_heading("Hook options (post 1)", 2)
para("A. I graded 114 live AI agents on security. 100 percent unsigned. 77 percent need no login. A thread on the A2A card problem.")
para("B. The fastest way to attack an AI agent is to read the file it publishes about itself.")
para("C. Valid is not safe. 114 live AI agents worked fine. Average security grade: C.")

doc.add_heading("Thread", 2)
para("1/ I graded the security posture of 114 live AI agents. Each publishes an agent card: a public file saying where it lives, how to authenticate, and what it can do. In the A2A protocol that card is the unit of trust. The results were not close.")
para("2/ All 114 were unsigned. 88 of 114 (about 3 in 4) required no authentication. Average grade: C. The few A's were unsigned too. They just had less else wrong.")
para("3/ Read a card as an attacker and it is a map. Endpoints show where to knock. The auth block shows whether anyone checks IDs. Push-notification webhooks can become a way to make the agent reach into a private network if it does not validate the target.")
para("4/ The unsigned result is the systemic one. The spec defines card signing over a canonical copy of the card. A signature proves the card you fetched is the card the operator published. Without it, anyone in the path can swap the endpoint or the auth rules.")
para("5/ That is a supply-chain weakness on the one file the whole protocol trusts. A demo can be open on purpose. No one ships a card they want forged.")
para("6/ Conformance tools (a2a-inspector, a2a-tck) say a card is valid. They do not say it is safe. So I built a2a-audit: reads the card like an attacker, maps issues to the OWASP agentic top 10, returns one grade you can fail a build on. Skill checks run on a local model, no key, no data leaves your box.")
para(f"7/ Grade any card yourself in the browser, all 114 results are baked in: {DEMO}  Open source: {REPO}")

rule()
doc.add_heading("Sources", 1)
para("OWASP Top 10 for Agentic Applications (2026), published 2025-12-09: https://genai.owasp.org/resource/owasp-top-10-for-agentic-applications-for-2026/")
para("A2A protocol specification (latest, v1.0): https://a2a-protocol.org/latest/specification/")
para("A2A card signing (spec section 8.4) + JSON Canonicalization Scheme (RFC 8785): https://www.rfc-editor.org/rfc/rfc8785.html")
para("Conformance / debugging tools (for contrast): https://github.com/a2aproject/a2a-inspector , https://github.com/a2aproject/a2a-tck")
para("Issue scanner (for contrast): https://github.com/cisco-ai-defense/a2a-scanner")
para(f"a2a-audit (this tool): {REPO}")
para("Posture stats: a2a-audit run over 114 live agent cards from a2aregistry.org and public A2A directories, captured 2026-05-29, canonical re-fetch, on-device DeBERTa classifier. Aggregate and anonymized; no individual agent named.")

doc.save(OUT)
print("wrote", OUT)
